import os, math, textwrap
from pathlib import Path
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = Path('/mnt/data/logistics_analysis_assets')
OUT.mkdir(exist_ok=True)
np.random.seed(42)

# ----------------------------
# 1. Simulate logistics data
# ----------------------------
n = 300
regions = ['North', 'South', 'East', 'West']
modes = ['Road', 'Rail', 'Air', 'Sea']
region = np.random.choice(regions, size=n, p=[0.27, 0.29, 0.21, 0.23])
mode = np.random.choice(modes, size=n, p=[0.55, 0.18, 0.12, 0.15])
dates = pd.date_range('2026-01-01', periods=90, freq='D')
date = np.random.choice(dates, size=n)
distance = np.clip(np.random.gamma(shape=3.1, scale=150, size=n) + 50, 40, 1800)
volume = np.clip(np.random.lognormal(mean=5.7, sigma=0.55, size=n), 30, 1800)

speed = {'Road': 48, 'Rail': 55, 'Air': 520, 'Sea': 32}
base = np.array([speed[m] for m in mode])
transit = distance / base
handling = np.array([3.5 if m=='Air' else 5.0 if m=='Rail' else 7.0 if m=='Road' else 12.0 for m in mode])
weather = np.random.gamma(1.5, 0.7, n)
congestion = np.where((mode=='Road') & (np.isin(region,['North','West'])), np.random.gamma(2.2, 1.2, n), np.random.gamma(1.0, 0.5, n))
expected = transit + handling
promised = np.maximum(expected + np.random.normal(1.6, 1.2, n), 2.5)
# Delay linked to road, distance, congestion and weather
mode_delay = np.where(mode=='Road', 1.9, np.where(mode=='Rail', 1.0, np.where(mode=='Air', 0.4, 1.2)))
delay = np.maximum(0, mode_delay + 0.004*distance + 0.45*weather + 0.65*congestion + np.random.normal(0,1.0,n) - 2.8)
delivery_time = promised + delay + np.random.normal(0,0.8,n)
delivery_time = np.maximum(delivery_time, 1.0)

# Cost model: fixed + distance + weight + mode factor
mode_cost_km = {'Road': 2.9, 'Rail': 1.8, 'Air': 12.0, 'Sea': 1.25}
mode_fixed = {'Road': 900, 'Rail': 1300, 'Air': 2500, 'Sea': 1700}
transport_cost = np.array([mode_fixed[m] for m in mode]) + np.array([mode_cost_km[m] for m in mode])*distance + 0.85*volume + np.random.normal(0, 420, n)
transport_cost = np.maximum(transport_cost, 500)
fuel_cost = transport_cost * np.where(mode=='Road', 0.24, np.where(mode=='Rail',0.16,np.where(mode=='Air',0.12,0.09)))
damage = np.clip(np.random.beta(1.6, 18, n) + np.where(mode=='Road',0.008,0), 0, 0.15)
on_time = delivery_time <= promised
efficiency = volume / delivery_time

# Introduce some missing values to demonstrate data-quality handling
raw = pd.DataFrame({
    'Shipment_ID': [f'SHP{1001+i}' for i in range(n)],
    'Date': pd.to_datetime(date),
    'Region': region,
    'Transport_Mode': mode,
    'Distance_km': distance,
    'Shipment_Volume_kg': volume,
    'Promised_Time_hr': promised,
    'Delivery_Time_hr': delivery_time,
    'Delay_Hours': delay,
    'Transport_Cost_INR': transport_cost,
    'Fuel_Cost_INR': fuel_cost,
    'Damage_Rate': damage,
    'On_Time': on_time,
    'Operational_Efficiency_kg_hr': efficiency,
})
raw.loc[np.random.choice(raw.index, 9, replace=False), 'Distance_km'] = np.nan
raw.loc[np.random.choice(raw.index, 7, replace=False), 'Shipment_Volume_kg'] = np.nan
raw.loc[np.random.choice(raw.index, 6, replace=False), 'Transport_Cost_INR'] = np.nan

# Clean
clean = raw.copy()
num_cols = ['Distance_km','Shipment_Volume_kg','Transport_Cost_INR']
for c in num_cols:
    clean[c] = clean[c].fillna(clean[c].median())
clean['Fuel_Cost_INR'] = clean['Fuel_Cost_INR'].fillna(clean['Fuel_Cost_INR'].median())

# ----------------------------
# 2. EDA statistics
# ----------------------------
desc = clean[['Distance_km','Shipment_Volume_kg','Promised_Time_hr','Delivery_Time_hr','Delay_Hours','Transport_Cost_INR','Fuel_Cost_INR','Damage_Rate','Operational_Efficiency_kg_hr']].describe().T
corr = clean[['Distance_km','Shipment_Volume_kg','Delivery_Time_hr','Delay_Hours','Transport_Cost_INR','Fuel_Cost_INR','Damage_Rate','Operational_Efficiency_kg_hr']].corr()
mode_summary = clean.groupby('Transport_Mode').agg(
    Shipments=('Shipment_ID','count'),
    Avg_Delivery_hr=('Delivery_Time_hr','mean'),
    Avg_Delay_hr=('Delay_Hours','mean'),
    On_Time_Rate=('On_Time','mean'),
    Avg_Cost_INR=('Transport_Cost_INR','mean'),
    Avg_Efficiency=('Operational_Efficiency_kg_hr','mean')
).reset_index()
region_summary = clean.groupby('Region').agg(
    Shipments=('Shipment_ID','count'),
    Avg_Delay_hr=('Delay_Hours','mean'),
    On_Time_Rate=('On_Time','mean'),
    Avg_Cost_INR=('Transport_Cost_INR','mean')
).reset_index()

# ----------------------------
# 3. Visualizations
# ----------------------------
plt.rcParams.update({'figure.dpi': 140, 'font.size': 10})

fig, ax = plt.subplots(figsize=(7.5,4.4))
ax.hist(clean['Delivery_Time_hr'], bins=24, edgecolor='black')
ax.axvline(clean['Delivery_Time_hr'].mean(), linestyle='--', linewidth=1.5, label=f"Mean = {clean['Delivery_Time_hr'].mean():.1f} hr")
ax.set_title('Distribution of Delivery Times')
ax.set_xlabel('Delivery time (hours)')
ax.set_ylabel('Number of shipments')
ax.legend()
fig.tight_layout(); fig.savefig(OUT/'01_delivery_distribution.png', bbox_inches='tight'); plt.close(fig)

fig, ax = plt.subplots(figsize=(7.5,4.4))
order = mode_summary.sort_values('Avg_Delay_hr', ascending=False)
ax.bar(order['Transport_Mode'], order['Avg_Delay_hr'])
ax.set_title('Average Delay by Transportation Mode')
ax.set_xlabel('Transportation mode'); ax.set_ylabel('Average delay (hours)')
fig.tight_layout(); fig.savefig(OUT/'02_delay_by_mode.png', bbox_inches='tight'); plt.close(fig)

fig, ax = plt.subplots(figsize=(7.5,4.4))
for m in modes:
    sub = clean[clean['Transport_Mode']==m]
    ax.scatter(sub['Distance_km'], sub['Transport_Cost_INR'], s=18, alpha=0.55, label=m)
ax.set_title('Transportation Distance vs Cost')
ax.set_xlabel('Distance (km)'); ax.set_ylabel('Transport cost (INR)')
ax.legend(ncol=2, fontsize=8)
fig.tight_layout(); fig.savefig(OUT/'03_distance_cost.png', bbox_inches='tight'); plt.close(fig)

fig, ax = plt.subplots(figsize=(7.5,4.4))
heat = corr.loc[['Distance_km','Shipment_Volume_kg','Delivery_Time_hr','Delay_Hours','Transport_Cost_INR','Fuel_Cost_INR','Damage_Rate','Operational_Efficiency_kg_hr'],
                ['Distance_km','Shipment_Volume_kg','Delivery_Time_hr','Delay_Hours','Transport_Cost_INR','Fuel_Cost_INR','Damage_Rate','Operational_Efficiency_kg_hr']]
im = ax.imshow(heat.values, vmin=-1, vmax=1, cmap='coolwarm', aspect='auto')
ax.set_xticks(range(len(heat.columns))); ax.set_xticklabels(heat.columns, rotation=45, ha='right', fontsize=8)
ax.set_yticks(range(len(heat.index))); ax.set_yticklabels(heat.index, fontsize=8)
for i in range(len(heat.index)):
    for j in range(len(heat.columns)):
        ax.text(j,i,f'{heat.iloc[i,j]:.2f}',ha='center',va='center',fontsize=7)
ax.set_title('Correlation Matrix of Key Logistics Metrics')
fig.colorbar(im, ax=ax, fraction=0.035, pad=0.03)
fig.tight_layout(); fig.savefig(OUT/'04_correlation_heatmap.png', bbox_inches='tight'); plt.close(fig)

fig, ax = plt.subplots(figsize=(7.5,4.4))
reg = region_summary.sort_values('On_Time_Rate')
ax.bar(reg['Region'], reg['On_Time_Rate']*100)
ax.set_title('On-Time Delivery Rate by Region')
ax.set_xlabel('Region'); ax.set_ylabel('On-time rate (%)')
ax.set_ylim(0,100)
for i,v in enumerate(reg['On_Time_Rate']*100): ax.text(i,v+2,f'{v:.1f}%',ha='center',fontsize=9)
fig.tight_layout(); fig.savefig(OUT/'05_region_ontime.png', bbox_inches='tight'); plt.close(fig)

fig, ax = plt.subplots(figsize=(7.5,4.4))
trend = clean.assign(Date=clean['Date']).groupby('Date').agg(Avg_Delay=('Delay_Hours','mean')).reset_index()
rolling = trend.set_index('Date')['Avg_Delay'].rolling(7, min_periods=1).mean()
ax.plot(trend['Date'], trend['Avg_Delay'], linewidth=1, alpha=0.45, label='Daily average')
ax.plot(rolling.index, rolling.values, linewidth=2, label='7-day rolling average')
ax.set_title('Daily Average Delay Trend')
ax.set_xlabel('Date'); ax.set_ylabel('Delay (hours)'); ax.legend()
fig.autofmt_xdate(); fig.tight_layout(); fig.savefig(OUT/'06_delay_trend.png', bbox_inches='tight'); plt.close(fig)

# ----------------------------
# 4. Generate DOCX
# ----------------------------
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7); sec.left_margin = Inches(0.75); sec.right_margin = Inches(0.75)

# styles
styles = doc.styles
styles['Normal'].font.name = 'Aptos'; styles['Normal'].font.size = Pt(10.5)
for sty_name, size in [('Title',24),('Heading 1',17),('Heading 2',13),('Heading 3',11)]:
    styles[sty_name].font.name='Aptos Display'; styles[sty_name].font.size=Pt(size)
    styles[sty_name].font.bold=True
    styles[sty_name].paragraph_format.keep_with_next = True

# helper functions
def shade_cell(cell, fill='D9EAF7'):
    tcPr = cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def set_cell_text(cell, text, bold=False):
    cell.text=''
    p=cell.paragraphs[0]; r=p.add_run(str(text)); r.bold=bold; r.font.size=Pt(9)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(df, title=None, percent_cols=None, money_cols=None, decimals=2):
    if title:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(title); r.bold=True; r.font.size=Pt(10)
    t=doc.add_table(rows=1, cols=len(df.columns)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for j,c in enumerate(df.columns): set_cell_text(t.rows[0].cells[j], c, True); shade_cell(t.rows[0].cells[j])
    for _,row in df.iterrows():
        cells=t.add_row().cells
        for j,c in enumerate(df.columns):
            val=row[c]
            if percent_cols and c in percent_cols: txt=f'{val*100:.1f}%'
            elif money_cols and c in money_cols: txt=f'₹{val:,.0f}'
            elif isinstance(val,float): txt=f'{val:.{decimals}f}'
            else: txt=str(val)
            set_cell_text(cells[j],txt)
    doc.add_paragraph()
    return t

def add_bullets(items):
    for x in items: doc.add_paragraph(x, style='List Bullet')

def add_code(text):
    for line in text.strip('\n').split('\n'):
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.25); p.paragraph_format.space_after=Pt(0)
        r=p.add_run(line); r.font.name='Consolas'; r.font.size=Pt(8.5)

def add_fig(path, caption):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(path), width=Inches(6.7))
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(caption); r.italic=True; r.font.size=Pt(9)

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.space_after=Pt(12)
r=p.add_run('ADVANCED DATA ANALYSIS AND VISUALIZATION IN LOGISTICS'); r.bold=True; r.font.size=Pt(23)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('A Python-Based Exploratory Analysis of a Simulated Logistics Dataset'); r.font.size=Pt(14)
for _ in range(5): doc.add_paragraph('')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Academic Project Report'); r.bold=True; r.font.size=Pt(13)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Prepared using simulated logistics operations data\nAugust 2026').font.size=Pt(11)
doc.add_page_break()

# Executive summary
h=doc.add_heading('Executive Summary', level=1)
summary=(f"This report demonstrates an end-to-end analytical workflow for logistics data using Python. A hypothetical dataset of {len(clean)} shipments was simulated across four regions and four transportation modes. The variables represent shipment distance, volume, promised and actual delivery time, delay, transport cost, fuel cost, damage rate, on-time status, and operational efficiency. The dataset intentionally includes missing observations so that data-quality treatment can be illustrated before analysis.\n\n"
         f"The exploratory analysis indicates that operational performance is influenced by transportation mode, shipment distance, and congestion-related delays. The average delivery time is {clean.Delivery_Time_hr.mean():.1f} hours, while the mean delay is {clean.Delay_Hours.mean():.1f} hours. Road transport shows the highest average delay at {mode_summary.loc[mode_summary.Transport_Mode=='Road','Avg_Delay_hr'].iloc[0]:.1f} hours, whereas air transport provides the fastest service. Transportation cost increases strongly with distance, making route length one of the main cost drivers. Regional comparison also reveals differences in on-time performance.\n\n"
         "Six visualizations were used to convert these findings into operationally useful evidence: a delivery-time distribution, delay-by-mode comparison, distance-versus-cost scatter plot, correlation heatmap, regional on-time chart, and time-series delay trend. The final recommendations focus on mode-specific planning, tighter monitoring of road shipments, route optimization, dynamic capacity allocation, and regular dashboard-based KPI monitoring.")
doc.add_paragraph(summary)

# 1 Introduction

doc.add_heading('1. Introduction', level=1)
doc.add_paragraph('Logistics operations generate data at nearly every stage of the supply chain, including order processing, transportation, warehousing, delivery, fuel consumption, and service-level compliance. Advanced analytics helps transform these records into measurable evidence for improving cost efficiency, delivery reliability, and resource utilization. The purpose of this project is to demonstrate how Python can be used to simulate, explore, visualize, and interpret logistics data in a structured manner.')
doc.add_paragraph('The analysis is designed as a hypothetical academic exercise. It does not represent the actual performance of a specific logistics company. Instead, the simulated values are constructed to resemble realistic operational relationships and to demonstrate the application of EDA techniques.')

# 2 Objectives

doc.add_heading('2. Objectives of the Analysis', level=1)
add_bullets([
    'Simulate a structured logistics dataset containing operational, cost, and service-level variables.',
    'Identify and treat missing values before statistical analysis.',
    'Calculate central tendency, dispersion, distributional summaries, and correlations.',
    'Create visualizations that reveal trends, relationships, and operational bottlenecks.',
    'Interpret findings in terms of delivery reliability, cost drivers, and efficiency.',
    'Develop practical recommendations for logistics decision-making based on the evidence.'
])

# 3 Dataset

doc.add_heading('3. Data Simulation and Dataset Structure', level=1)
doc.add_paragraph(f'A total of {n} shipment records were generated for the period 1 January 2026 to 31 March 2026. The dataset covers four regions (North, South, East, West) and four transport modes (Road, Rail, Air, Sea). The simulation uses probabilistic distributions so that values vary across shipments while preserving plausible logistics relationships.')
fields = pd.DataFrame([
['Shipment_ID','Unique identifier for each shipment'],['Date','Shipment or delivery date'],['Region','Operating region'],['Transport_Mode','Primary mode used for transport'],['Distance_km','Shipment distance in kilometres'],['Shipment_Volume_kg','Shipment weight/volume proxy in kilograms'],['Promised_Time_hr','Customer or SLA promised delivery time'],['Delivery_Time_hr','Observed delivery time'],['Delay_Hours','Hours beyond planned operational timing'],['Transport_Cost_INR','Estimated transportation cost'],['Fuel_Cost_INR','Estimated fuel-related cost'],['Damage_Rate','Proportion of damaged shipment value/units'],['On_Time','Whether delivery met promised time'],['Operational_Efficiency_kg_hr','Shipment volume handled per delivery hour']
], columns=['Variable','Description'])
add_table(fields, 'Table 1. Dataset variables')

# 4 Data quality

doc.add_heading('4. Data Cleaning and Preprocessing', level=1)
doc.add_paragraph('To demonstrate a practical preprocessing pipeline, missing values were intentionally inserted into Distance_km, Shipment_Volume_kg, and Transport_Cost_INR. The missing observations were replaced using the median of each variable. Median imputation is appropriate for this exercise because logistics measures such as distance, shipment volume, and cost may contain skewed distributions and outliers; the median is less sensitive to extreme values than the mean.')
missing = raw.isna().sum().reset_index(); missing.columns=['Variable','Missing_Count']; missing=missing[missing.Missing_Count>0]
add_table(missing, 'Table 2. Introduced missing values before cleaning')
cleaning_code = """
num_cols = ['Distance_km', 'Shipment_Volume_kg', 'Transport_Cost_INR']
for col in num_cols:
    df[col] = df[col].fillna(df[col].median())
"""
doc.add_heading('Python preprocessing excerpt', level=2); add_code(cleaning_code)

# 5 EDA

doc.add_heading('5. Exploratory Data Analysis', level=1)
doc.add_heading('5.1 Descriptive Statistics', level=2)
stat_view=desc[['count','mean','50%','std','min','max']].copy().reset_index().rename(columns={'index':'Metric','50%':'Median','std':'Std. Dev.'})
add_table(stat_view, 'Table 3. Descriptive statistics')
doc.add_paragraph(f"The average transport cost is approximately ₹{clean.Transport_Cost_INR.mean():,.0f}, while the median is ₹{clean.Transport_Cost_INR.median():,.0f}. The mean delivery time of {clean.Delivery_Time_hr.mean():.1f} hours is above the median of {clean.Delivery_Time_hr.median():.1f} hours, indicating some longer deliveries are pulling the mean upward. Delay is also positively skewed because most shipments have modest delays while a smaller group experiences more severe disruptions.")

doc.add_heading('5.2 Transportation Mode Performance', level=2)
add_table(mode_summary, 'Table 4. Mode-level logistics performance', percent_cols=['On_Time_Rate'])
road_delay=float(mode_summary.loc[mode_summary.Transport_Mode=='Road','Avg_Delay_hr'].iloc[0]); air_time=float(mode_summary.loc[mode_summary.Transport_Mode=='Air','Avg_Delivery_hr'].iloc[0])
doc.add_paragraph(f'Road has the highest average delay ({road_delay:.1f} hours), making it the most important candidate for operational improvement. Air provides the shortest average delivery time at approximately {air_time:.1f} hours, but its higher unit transportation cost makes it unsuitable as a blanket replacement for lower-cost modes. The analytical implication is that transport mode should be selected by service requirement rather than speed alone.')

# 6 Visualizations

doc.add_heading('6. Visual Analysis and Interpretation', level=1)
doc.add_heading('6.1 Delivery-Time Distribution', level=2)
add_fig(OUT/'01_delivery_distribution.png','Figure 1. Distribution of delivery times')
doc.add_paragraph('A histogram is used because it reveals the shape of the delivery-time distribution, including concentration, spread, and potential long-tail behaviour. The distribution shows that most shipments are clustered in a practical operating range, while a smaller number take substantially longer. Such a long tail is operationally important because extreme delays can increase customer complaints, re-delivery effort, and resource requirements.')

doc.add_page_break()
doc.add_heading('6.2 Average Delay by Transportation Mode', level=2)
add_fig(OUT/'02_delay_by_mode.png','Figure 2. Average delay by transport mode')
doc.add_paragraph('A bar chart is appropriate for comparing categorical groups. The chart makes the difference in delay performance between transport modes immediately visible. Road shipments demonstrate the greatest delay exposure in this simulated dataset. Management should therefore investigate road-route congestion, loading delays, dispatch planning, and carrier adherence to scheduled departure times.')

doc.add_page_break()
doc.add_heading('6.3 Distance and Transportation Cost', level=2)
add_fig(OUT/'03_distance_cost.png','Figure 3. Relationship between distance and transportation cost')
cost_corr=float(corr.loc['Distance_km','Transport_Cost_INR'])
doc.add_paragraph(f'A scatter plot is suitable for examining the relationship between two continuous variables. The pattern shows a positive relationship between distance and transport cost, with a correlation of approximately {cost_corr:.2f}. The upward pattern indicates that longer routes generally require greater spending. The spread between modes also shows that the same distance can generate very different costs depending on the mode. This supports route consolidation and mode optimization as potential cost-control levers.')

doc.add_page_break()
doc.add_heading('6.4 Correlation Analysis', level=2)
add_fig(OUT/'04_correlation_heatmap.png','Figure 4. Correlation heatmap of key logistics variables')
doc.add_paragraph('A correlation heatmap provides a compact view of linear relationships among quantitative variables. The strongest business-relevant pattern is the positive association between distance and transportation cost, while delivery time and delay also move together. Correlation does not establish causation, so these relationships should be followed by operational investigation or predictive modelling before policy changes are implemented.')

doc.add_page_break()
doc.add_heading('6.5 On-Time Delivery by Region', level=2)
add_fig(OUT/'05_region_ontime.png','Figure 5. On-time delivery performance by region')
best=region_summary.loc[region_summary.On_Time_Rate.idxmax()]; worst=region_summary.loc[region_summary.On_Time_Rate.idxmin()]
doc.add_paragraph(f'The regional bar chart compares service-level performance. In the simulation, {best.Region} records the highest on-time rate at {best.On_Time_Rate*100:.1f}%, while {worst.Region} records the lowest at {worst.On_Time_Rate*100:.1f}%. Regional variation suggests that geography, congestion, warehouse hand-off quality, or carrier allocation can influence service reliability. A logistics manager should therefore avoid relying only on company-wide averages and should monitor regional KPIs separately.')

doc.add_page_break()
doc.add_heading('6.6 Daily Delay Trend', level=2)
add_fig(OUT/'06_delay_trend.png','Figure 6. Daily average delay and seven-day rolling average')
doc.add_paragraph('The line chart is used to observe changes over time. The seven-day rolling average reduces day-to-day noise and highlights persistent changes in delay performance. Peaks can indicate operational disruptions or demand pressure. In a real logistics environment, the trend should be enriched with weather, holiday, traffic, warehouse utilization, and carrier data to identify the underlying cause of spikes.')

# 7 Key insights

doc.add_heading('7. Key Analytical Insights', level=1)
insights=[
    f'Road transport is the main delay-risk area, with an average delay of about {road_delay:.1f} hours in the simulated data.',
    f'Distance is a major cost driver; its correlation with transport cost is approximately {cost_corr:.2f}.',
    f'Air transport achieves very short delivery times ({air_time:.1f} hours on average) but should be reserved for time-sensitive shipments because of its cost profile.',
    f'Regional service performance is not uniform: the gap between the best and worst on-time regions is {(best.On_Time_Rate-worst.On_Time_Rate)*100:.1f} percentage points.',
    'The delivery-time distribution and daily delay trend show why averages alone are insufficient; the long tail and temporary peaks need separate monitoring.',
    'Operational efficiency should be evaluated jointly with service and cost metrics rather than optimizing a single KPI.'
]
add_bullets(insights)

# 8 Recommendations

doc.add_heading('8. Recommendations for Logistics Decision-Making', level=1)
recs=[
('1. Improve road-route planning','Use route-level congestion monitoring, scheduled departure controls, and carrier performance scorecards for road shipments. Focus first on lanes with repeated high delay.'),
('2. Apply mode segmentation','Use air transport for urgent or high-value shipments, road for flexible regional moves, and rail/sea where lead-time tolerance allows lower-cost movement.'),
('3. Optimize route distance and consolidation','Combine compatible shipments, redesign delivery zones, and evaluate cross-docking or hub locations to reduce unnecessary kilometres.'),
('4. Monitor region-specific KPIs','Track on-time rate, average delay, cost per shipment, cost per kilometre, and damage rate by region and carrier.'),
('5. Build a logistics dashboard','Refresh operational KPIs daily or weekly so that managers can detect deteriorating performance before service levels materially decline.'),
('6. Extend the analysis to predictive models','Use regression or machine-learning models to estimate delay probability and expected cost using distance, mode, region, weather, demand, and historical carrier performance.')
]
for title,body in recs:
    p=doc.add_paragraph(); r=p.add_run(title+': '); r.bold=True; p.add_run(body)

# 9 Methodology

doc.add_heading('9. Methodology and Python Workflow', level=1)
doc.add_paragraph('The workflow follows six stages: data simulation, data-quality checking, cleaning, descriptive statistics, visualization, and interpretation. NumPy and pandas are used for numerical generation and tabular analysis, while Matplotlib is used for the visualizations. The process can be adapted directly to real CSV or Excel logistics data by replacing the simulation step with a data-import statement.')
code = """
import pandas as pd
import matplotlib.pyplot as plt

# Load real data in place of the simulated dataframe
# df = pd.read_csv('logistics_data.csv')

# Descriptive statistics
print(df.describe())

# Correlation matrix
corr = df[['Distance_km','Delivery_Time_hr','Delay_Hours',
           'Transport_Cost_INR']].corr()

# Mode comparison
mode_summary = df.groupby('Transport_Mode').agg(
    Avg_Delay=('Delay_Hours','mean'),
    Avg_Cost=('Transport_Cost_INR','mean'),
    On_Time_Rate=('On_Time','mean')
)

# Visualization example
plt.hist(df['Delivery_Time_hr'], bins=20, edgecolor='black')
plt.title('Distribution of Delivery Times')
plt.xlabel('Delivery time (hours)')
plt.ylabel('Shipments')
plt.show()
"""
add_code(code)

# 10 Limitations

doc.add_heading('10. Limitations of the Analysis', level=1)
add_bullets([
    'The dataset is simulated and therefore cannot be treated as evidence of the performance of a real company.',
    'The model does not include warehouse capacity, inventory levels, traffic feeds, weather APIs, carrier identifiers, or customer priority classes.',
    'Median imputation is a simple educational technique; production datasets should use a missing-data strategy based on the cause and pattern of missingness.',
    'Correlation measures association, not causation. Operational experiments or predictive modelling are required to validate causal assumptions.',
    'Cost figures are illustrative and exclude some real-world elements such as tolls, labour, insurance, taxes, and penalties.'
])

# 11 Conclusion

doc.add_heading('11. Conclusion', level=1)
doc.add_paragraph('This project demonstrates how advanced data analysis and visualization can support logistics management. The simulated dataset was first inspected and cleaned, after which descriptive statistics and correlation analysis were used to identify major patterns. Visualizations then converted those patterns into operational insights. The analysis points to transportation mode, distance, regional variation, and recurring delay events as important areas for managerial attention. Most importantly, the project shows that logistics decisions should combine service, cost, and efficiency measures rather than relying on a single KPI. A similar Python workflow can be applied to real shipment records to support continuous performance monitoring, route optimization, carrier evaluation, and predictive logistics planning.')

# 12 Appendix

doc.add_page_break()
doc.add_heading('12. Appendix: Sample Simulated Records', level=1)
sample=clean[['Shipment_ID','Date','Region','Transport_Mode','Distance_km','Shipment_Volume_kg','Delivery_Time_hr','Delay_Hours','Transport_Cost_INR','On_Time']].head(10).copy()
sample['Date']=sample['Date'].dt.strftime('%d-%b-%Y')
add_table(sample.round({'Distance_km':0,'Shipment_Volume_kg':0,'Delivery_Time_hr':1,'Delay_Hours':1,'Transport_Cost_INR':0}), 'Table 5. Sample cleaned shipment records', money_cols=['Transport_Cost_INR'])

# Header/footer
for section in doc.sections:
    header=section.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.CENTER; header.text='Advanced Data Analysis and Visualization in Logistics'
    for r in header.runs: r.font.size=Pt(8); r.italic=True
    footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=footer.add_run('Academic Project Report | Simulated Dataset'); rr.font.size=Pt(8)

# save
out_doc='/mnt/data/Advanced_Logistics_Data_Analysis_Report.docx'
doc.save(out_doc)
# Also save cleaned dataset and the script itself
clean.to_csv('/mnt/data/logistics_simulated_cleaned.csv', index=False)
print(out_doc)
print('assets:', OUT)
print('mean delivery', clean.Delivery_Time_hr.mean())
print('mean delay', clean.Delay_Hours.mean())
print('cost corr', cost_corr)
